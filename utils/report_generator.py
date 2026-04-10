"""
utils/report_generator.py
HostTrace AI — Report Engine v5.0

Architecture:
  1. generate_text_report(data)  → structured Markdown string  (pure text, no I/O)
  2. render_pdf_from_text(text, data) → PDF bytes via ReportLab  (layout only)
  3. generate_pdf_report(data)   → convenience wrapper (calls 1 then 2)

This separation means the same textual report can be:
  - Displayed in the browser   (raw Markdown)
  - Downloaded as a .txt file  (identical string)
  - Rendered into a styled PDF  (via ReportLab)
  - Passed to any future AI summariser or email template
"""

import os
import re
import textwrap
from io import BytesIO
import datetime
import xml.sax.saxutils as _saxutils

# ── ReportLab ────────────────────────────────────────────────────
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    HRFlowable, Table, TableStyle, Image, PageBreak,
)
from reportlab.pdfbase import pdfmetrics

# ── Matplotlib  (headless, for embedded charts) ──────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

# ── Font registration ────────────────────────────────────────────
try:
    from reportlab.pdfbase.ttfonts import TTFont
    _BASE = os.path.dirname(os.path.abspath(__file__))
    _FONT = os.path.join(_BASE, "DejaVuSans.ttf")
    if os.path.exists(_FONT):
        pdfmetrics.registerFont(TTFont("DejaVuSans", _FONT))
        FONT_NAME = "DejaVuSans"
    else:
        FONT_NAME = "Helvetica"
except Exception:
    FONT_NAME = "Helvetica"


# ══════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════

def _safe(text) -> str:
    """
    Convert any value to a plain ASCII string safe for Helvetica-mode PDF.
    Uses encode/decode instead of ord() to handle all Unicode safely.
    """
    s = str(text)
    if FONT_NAME == "Helvetica":
        # Replace common unicode punctuation with ASCII equivalents
        s = s.replace("\u2014", "-").replace("\u2013", "-")   # em/en dash
        s = s.replace("\u2018", "'").replace("\u2019", "'")   # curly quotes
        s = s.replace("\u201c", '"').replace("\u201d", '"')   # curly double quotes
        s = s.replace("\u2022", "*").replace("\u2026", "...") # bullet, ellipsis
        s = s.replace("\u2264", "<=").replace("\u2265", ">=") # comparison
        s = s.replace("\u2550", "=").replace("\u2500", "-")   # box-drawing
        # Strip remaining non-ASCII safely (no ord() call)
        s = s.encode("ascii", "ignore").decode("ascii")
    return s


SPECIAL_CHARS = {
    "\u2014": "-", "\u2013": "-",
    "\u2018": "'", "\u2019": "'",
    "\u201c": '"', "\u201d": '"',
    "\u2022": "*", "\u2026": "...",
    "\u2264": "<=", "\u2265": ">=",
    "\u2550": "=", "\u2500": "-",
    "\u25cf": "*", "\u2192": "->",
    # Status symbols -> ASCII
    "\u2713": "[OK]", "\u2717": "[X]",
    "\u26a0": "[!]",  "\u2714": "[OK]",
    # Emoji fallbacks  (most common in threat reports)
    "\U0001f534": "[HIGH]",  # red circle
    "\U0001f7e1": "[MED]",   # yellow circle
    "\U0001f7e2": "[LOW]",   # green circle
    "\U0001f6a8": "[ALERT]", # siren
    "\u2705": "[OK]",        # check mark
    "\u274c": "[X]",         # cross mark
}


def _p(text) -> str:
    """
    Convert text to a ReportLab Paragraph-safe string.
    Steps:
      1. Stringify and replace known Unicode with ASCII equivalents
      2. Strip remaining non-ASCII (no ord())
      3. Escape XML special chars (&, <, >) so ReportLab's XML parser doesn't crash
      4. Convert Markdown **bold** -> <b>bold</b>  and `code` -> <font ...>code</font>
    """
    s = str(text)
    for uni, asc in SPECIAL_CHARS.items():
        s = s.replace(uni, asc)
    # Strip any remaining non-ASCII safely
    s = s.encode("ascii", "ignore").decode("ascii")
    # Escape XML special characters BEFORE adding any markup tags
    s = _saxutils.escape(s)  # & -> &amp;  < -> &lt;  > -> &gt;
    # Now convert markdown bold:  **text** -> <b>text</b>
    s = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)
    # Inline code: `text` -> monospace
    s = re.sub(r"`([^`]+)`", r'<font face="Courier" fontSize="8">\1</font>', s)
    return s


def _divider(char: str = "─", width: int = 72) -> str:
    return char * width


def _section(title: str) -> str:
    bar = "═" * 72
    return f"\n{bar}\n  {title.upper()}\n{bar}\n"


# ══════════════════════════════════════════════════════════════════
# CHART GENERATORS  (return BytesIO PNG buffers)
# ══════════════════════════════════════════════════════════════════

def _chart_risk_trend(scan_history: list):
    """Line chart: risk score over scanned sessions."""
    if not scan_history or len(scan_history) < 2:
        return None
    try:
        labels = [f"#{i+1}\n{h.get('domain', '')[:12]}" for i, h in enumerate(scan_history)]
        values = [h.get("risk", 0) for h in scan_history]

        fig, ax = plt.subplots(figsize=(7, 3))
        ax.fill_between(range(len(values)), values, alpha=0.15, color="#2a5e82")
        ax.plot(range(len(values)), values, marker="o", color="#2a5e82",
                linewidth=2, markersize=6)
        for i, v in enumerate(values):
            ax.annotate(str(v), (i, v), textcoords="offset points",
                        xytext=(0, 8), ha="center", fontsize=8, color="#0f3d5e")
        ax.set_xticks(range(len(labels)))
        ax.set_xticklabels(labels, fontsize=7)
        ax.set_ylim(0, 105)
        ax.set_ylabel("Risk Score", fontsize=9)
        ax.set_title("Risk Trend History  (session scans)", fontsize=11, color="#0f3d5e", pad=10)
        ax.axhline(70, color="#ff4f6d", linestyle="--", linewidth=0.8, label="High threshold")
        ax.axhline(30, color="#ffd166", linestyle="--", linewidth=0.8, label="Medium threshold")
        ax.legend(fontsize=7)
        ax.grid(True, linestyle="--", alpha=0.4)
        ax.spines[["top", "right"]].set_visible(False)

        buf = BytesIO()
        plt.tight_layout()
        plt.savefig(buf, format="png", dpi=160)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception as e:
        print(f"[Report] Risk trend chart error: {e}")
        return None


def _chart_redirect_chain(chain: list):
    """Vertical flowchart for redirect hops."""
    if not chain or len(chain) <= 1:
        return None
    try:
        n = len(chain)
        fig, ax = plt.subplots(figsize=(7, max(2.5, n * 1.1)))
        ax.axis("off")
        ax.set_xlim(0, 1)
        ax.set_ylim(-0.2, n + 0.4)

        STATUS_COLORS = {
            "2": "#d1fae5",    # 2xx green
            "3": "#fef08a",    # 3xx yellow / redirect
            "4": "#fecaca",    # 4xx red
            "5": "#fecaca",    # 5xx red
        }

        for i, hop in enumerate(chain):
            y = n - i - 1
            status = hop.get("status", "?")
            url = hop.get("url", "—")[:65]
            label = f"Hop {hop.get('hop', i+1)}  [{status}]\n{url}"
            fc = STATUS_COLORS.get(str(status)[0], "#f1f5f9")
            final = hop.get("final", False)
            ec = "#22c55e" if final else "#94a3b8"
            lw = 2 if final else 1
            ax.text(0.5, y + 0.5, label, ha="center", va="center", fontsize=8,
                    bbox=dict(boxstyle="round,pad=0.4", fc=fc, ec=ec, lw=lw),
                    fontfamily="monospace")
            if i < n - 1:
                ax.annotate("", xy=(0.5, y + 0.05), xytext=(0.5, y + 0.45),
                            arrowprops=dict(arrowstyle="->", lw=1.4, color="#64748b"))

        ax.set_title("Redirect Chain Flow", fontsize=11, color="#0f3d5e", pad=8)
        buf = BytesIO()
        plt.tight_layout()
        plt.savefig(buf, format="png", dpi=160)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception as e:
        print(f"[Report] Redirect chart error: {e}")
        return None


def _chart_risk_gauge(risk_score: int, risk_level: str):
    """Semicircular gauge for risk score."""
    try:
        fig, ax = plt.subplots(figsize=(4, 2.2), subplot_kw=dict(aspect="equal"))
        ax.axis("off")

        def draw_arc(ax, theta1, theta2, color, radius=1, width=0.22):
            from matplotlib.patches import Wedge
            w = Wedge((0, 0), radius, theta1, theta2,
                      width=width, facecolor=color, edgecolor="white", lw=0.5)
            ax.add_patch(w)

        draw_arc(ax, 0, 60,   "#ff4f6d")   # high
        draw_arc(ax, 60, 120, "#ffd166")   # medium
        draw_arc(ax, 120, 180,"#00c97a")   # low

        angle = 180 - (risk_score / 100) * 180
        import math
        r = 0.85
        x = r * math.cos(math.radians(angle))
        y = r * math.sin(math.radians(angle))
        ax.annotate("", xy=(x, y), xytext=(0, 0),
                    arrowprops=dict(arrowstyle="-|>", lw=2, color="#1e3a5f"),
                    annotation_clip=False)
        ax.set_xlim(-1.2, 1.2)
        ax.set_ylim(-0.3, 1.2)
        ax.text(0, -0.15, f"{risk_score}", ha="center", va="center",
                fontsize=22, fontweight="bold", color="#1e3a5f")
        ax.text(0, -0.3, f"Risk Score  /  {risk_level}", ha="center",
                fontsize=9, color="#4a7a9b")

        buf = BytesIO()
        plt.tight_layout()
        plt.savefig(buf, format="png", dpi=160)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception as e:
        print(f"[Report] Gauge chart error: {e}")
        return None


# ══════════════════════════════════════════════════════════════════
# STEP 1 — STRUCTURED TEXT / MARKDOWN REPORT
# ══════════════════════════════════════════════════════════════════

def generate_text_report(data: dict) -> str:
    """
    Builds a fully self-contained, human-readable Markdown report from scan data.
    This pure-text output is the canonical source of truth for everything that
    follows (PDF, email, display, AI summarisation, etc.).
    """
    lines = []
    add = lines.append

    scan_dt = data.get("scan_timestamp",
                        datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC"))
    domain  = data.get("domain", "Unknown")
    input_url = data.get("input_url", domain)

    # ── Cover ───────────────────────────────────────────────────
    add("# HostTrace AI — Forensic Intelligence Report")
    add("")
    add(_divider("═"))
    add(f"**Scan Timestamp :** {scan_dt}")
    add(f"**Report ID      :** {data.get('report_id', 'N/A')}")
    add(f"**Trace ID       :** {data.get('trace_id', 'N/A')}")
    add(f"**Input URL/IP   :** {input_url}")
    add(f"**Canonical Domain:** {domain}")
    add(f"**IP Address(es) :** {', '.join(data.get('ip_addresses', [])) or 'Unresolved'}")
    add(_divider("═"))
    add("")

    # ── Section 1: Overview ──────────────────────────────────────
    add(_section("1. Overview"))
    risk_score = data.get("risk_score", 0)
    risk_level = data.get("risk_level", "Unknown")
    verdict    = data.get("verdict", {})
    ai_conf    = data.get("ai_confidence", {})
    pred       = data.get("ai_prediction", {})

    if risk_score >= 71:
        severity, severity_sym = "HIGH", "🔴"
    elif risk_score >= 31:
        severity, severity_sym = "MEDIUM", "🟡"
    else:
        severity, severity_sym = "LOW", "🟢"

    add(f"**Final Status   :** {verdict.get('status', '—')}")
    add(f"**Risk Score     :** {risk_score} / 100  →  {severity_sym} {severity}")
    add(f"**Risk Level     :** {risk_level}")
    add(f"**AI Confidence  :** {verdict.get('confidence', ai_conf.get('ai_confidence_pct', '—'))}%")
    add(f"**Infrastructure :** {data.get('real_hosting', '—')}")
    proxy_txt  = (f"Detected ({data.get('proxy_provider','?')})"
                  if data.get("proxy_detected") else "Not Detected")
    add(f"**Proxy / CDN    :** {proxy_txt}")
    add(f"**Masking Level  :** {data.get('masking_level', '—')}")
    add("")

    # ML Classification block
    if pred:
        add("### AI Model Classification")
        add(f"- **Model :** RandomForestClassifier (local, offline)  |  Accuracy: {pred.get('accuracy', 91.6):.1f}%")
        add(f"- **Prediction :** `{pred.get('prediction', 'UNKNOWN')}`")
        probs = pred.get("probabilities", {})
        add(f"- **Probability matrix :**  "
            f"SAFE={probs.get('SAFE', 0):.1f}%  |  "
            f"SUSPICIOUS={probs.get('SUSPICIOUS', 0):.1f}%  |  "
            f"DANGEROUS={probs.get('DANGEROUS', 0):.1f}%")
        add("")

    # Why risky explanation
    risk_reason = data.get("risk_reason", {})
    if risk_reason and isinstance(risk_reason, dict):
        add("### Risk Explanation (XAI — Local Rule Engine)")
        add(f"**{risk_reason.get('title', '')}**")
        for pt in risk_reason.get("points", []):
            add(f"  - {pt}")
        add("")

    # ── Section 2: OSINT Analysis ────────────────────────────────
    add(_section("2. OSINT Analysis"))

    # WHOIS
    add("### 2.1  WHOIS Forensics")
    whois = data.get("whois", {})
    for k, label in [
        ("registrar",     "Registrar"),
        ("org",           "Organisation"),
        ("country",       "Country"),
        ("creation_date", "Created"),
        ("expiry_date",   "Expires"),
        ("dnssec",        "DNSSEC"),
    ]:
        add(f"- **{label} :** {whois.get(k, '—')}")
    nss = whois.get("name_servers", [])
    add(f"- **Name Servers :** {', '.join(nss) if nss else '—'}")
    add("")

    # DNS / IP
    add("### 2.2  DNS Resolution")
    add(f"- **Resolved :** {'Yes' if data.get('ip_addresses') else 'No'}")
    add(f"- **TTL Hint :** {data.get('ttl_hint', '—')}")
    add(f"- **IPs :** {', '.join(data.get('ip_addresses', ['—']))}")
    add("")

    # Hosting
    add("### 2.3  Hosting & CDN Detection")
    add(f"- **Detected Hosting :** {data.get('real_hosting', '—')}")
    add(f"- **Possible Hosting :** {data.get('possible_hosting') or '—'}")
    add(f"- **Hosting Confidence :** {data.get('confidence', '—')}%")
    add(f"- **CDN Detected :** {'Yes — ' + str(data.get('cdn_provider','')) if data.get('cdn_detected') else 'No'}")
    add(f"- **WAF Suspected :** {'Yes' if data.get('waf_suspected') else 'No'}")
    add("")

    # OSINT Simulation
    osint = data.get("osint_simulation", {})
    if osint:
        add("### 2.4  Historical Domain Reputation")
        add(f"- {osint.get('exposure_note', '—')}")
        add(f"- **Alt Providers (historical) :** {', '.join(osint.get('alt_providers', [])) or '—'}")
        recs = osint.get("records", [])
        if recs:
            add("")
            add("| Date | IP | Provider | Record Type |")
            add("|------|----|----------|-------------|")
            for r in recs:
                add(f"| {r.get('date','?')} | {r.get('ip','?')} | "
                    f"{r.get('provider','?')} | {r.get('type','?')} |")
        add("")

    # ── Section 3: Threat Intelligence ───────────────────────────
    add(_section("3. Threat Intelligence"))

    threat = data.get("threat_intel", {})
    add(f"- **VirusTotal Flags :** {threat.get('virustotal_flags', 0)}")
    add(f"- **Blacklist Hits   :** {threat.get('blacklist_hits', 0)}")
    add(f"- **Abuse.ch C2      :** {'Yes ⚠' if threat.get('abuse_ch') else 'No'}")
    add(f"- **OTX Pulses       :** {threat.get('otx_pulses', 0)}")
    add(f"- **C2 Detected      :** {'Yes ⚠' if threat.get('c2_detected') else 'No'}")
    add(f"- **Phishing Category:** {'Yes ⚠' if threat.get('phishing_category') else 'No'}")
    add("")

    # SSL/TLS
    add("### 3.1  SSL / TLS Certificate")
    ssl = data.get("ssl_analysis", {})
    add(f"- **Valid        :** {'Yes ✓' if ssl.get('ssl_valid') else 'No ✗'}")
    add(f"- **Expired      :** {'Yes ⚠' if ssl.get('ssl_expired') else 'No'}")
    add(f"- **Self-Signed  :** {'Yes ⚠' if ssl.get('self_signed') else 'No'}")
    add(f"- **Grade        :** {ssl.get('grade', 'N/A')}")
    add(f"- **Issuer       :** {ssl.get('issuer', {}).get('organizationName', ssl.get('issuer', '—'))}")
    add(f"- **Valid From   :** {ssl.get('valid_from', '—')}")
    add(f"- **Valid Until  :** {ssl.get('valid_until', '—')}")
    add(f"- **Days Left    :** {ssl.get('days_remaining', '—')}")
    add("")

    # HTTP Security Headers
    add("### 3.2  HTTP Security Header Analysis")
    http = data.get("http_analysis", {})
    add(f"- **Server        :** {http.get('server', 'Unknown')}")
    add(f"- **Security Score:** {http.get('security_score', 0)}%")
    add(f"- **HSTS          :** {'Enabled ✓' if http.get('hsts') else 'Missing ✗'}")
    add(f"- **CSP           :** {'Present ✓' if http.get('csp') else 'Missing ✗'}")
    add(f"- **X-Frame-Opts  :** {'Present ✓' if http.get('x_frame') else 'Missing ✗'}")
    present = http.get("present_security_headers", [])
    missing = http.get("missing_security_headers", [])
    if present:
        add(f"- **Present Headers:** {', '.join(present)}")
    if missing:
        add(f"- **Missing Headers :** {', '.join(missing)}")
    add("")

    # Hidden Proxy / Cloaking
    add("### 3.3  Proxy & Cloaking Detection")
    origin = data.get("origin_discovery", {})
    asn    = data.get("asn_analysis", {})
    add(f"- **Proxy Active     :** {'Yes — ' + str(data.get('proxy_provider','?')) if data.get('proxy_detected') else 'No'}")
    add(f"- **Masking Level    :** {data.get('masking_level', '—')}")
    add(f"- **Origin IPs Found :** {', '.join(origin.get('possible_origin_ips', [])) or 'None'}")
    add(f"- **Subdomains Scanned:** {origin.get('subdomains_checked', 0)}")
    add(f"- **ASN Mismatch     :** {'Yes ⚠ — ' + asn.get('mismatch_note','') if asn.get('mismatch_detected') else 'No'}")
    add(f"- **Proxy Layer ASN  :** {asn.get('proxy_provider', '—')}")
    add(f"- **Origin ASN       :** {asn.get('origin_asn_provider', '—')}")
    add("")

    # ── Section 4: Redirect Chain ─────────────────────────────────
    add(_section("4. Redirect Chain Analysis"))
    rc = data.get("redirect_chain", {})
    add(f"- **Total Hops     :** {rc.get('total_hops', 0)}")
    add(f"- **Redirect Count :** {rc.get('redirect_count', 0)}")
    suspicious_chain = rc.get("suspicious", False)
    add(f"- **Suspicious     :** {'Yes ⚠ — Long redirect chain detected' if suspicious_chain else 'No'}")
    add("")

    chain = rc.get("chain", [])
    if chain:
        add("**Hop-by-Hop Table:**")
        add("")
        add("| Hop | Status | URL | Final |")
        add("|-----|--------|-----|-------|")
        for hop in chain:
            url_trunc = str(hop.get("url", "—"))[:70]
            final_tag = "✓ FINAL" if hop.get("final") else "→ redirect"
            add(f"| {hop.get('hop','?')} | {hop.get('status','?')} "
                f"| {url_trunc} | {final_tag} |")
        add("")
        add("> **[Chart]** A visual flowchart of the redirect chain is embedded in the PDF below this section.")
    else:
        add("No redirect hops recorded. Direct connection established.")
    add("")

    # ── Section 5: AI Risk Prediction ────────────────────────────
    add(_section("5. AI Risk Prediction"))

    add(f"### Final Verdict : `{verdict.get('status', '—')}`")
    add(f"- **Risk Score     :** {risk_score} / 100  ({severity} severity)")
    add(f"- **Confidence     :** {verdict.get('confidence', '—')}%")
    add(f"- **AI Pattern     :** {ai_conf.get('infrastructure_pattern', '—')}")
    add(f"- **AI Summary     :** {ai_conf.get('summary', '—')}")
    add("")

    # Contributing factors
    factors = data.get("threat_flags", data.get("explanation", []))
    if factors:
        add("### Key Contributing Factors")
        for f in factors[:12]:
            add(f"  - {f}")
        add("")

    # Risk breakdown
    rb = data.get("risk_breakdown", {})
    if rb:
        add("### Risk Factor Breakdown")
        add("")
        add("| Component | Score Impact |")
        add("|-----------|-------------|")
        for k, v in rb.items():
            if v != 0:
                sign = "+" if v > 0 else ""
                add(f"| {k.replace('_',' ').title()} | {sign}{v} |")
        add("")

    # XAI signals
    xai = data.get("xai_signals", [])
    if xai:
        add("### Explainable AI (XAI) Feature Contributions")
        for sig in xai:
            add(f"  - **{sig.get('feature','?')}** [{sig.get('impact','?')}]: {sig.get('desc','')}")
        add("")

    # ── Section 6: Geolocation & ASN ─────────────────────────────
    add(_section("6. Geolocation & ASN Intelligence"))
    geo = data.get("geo_analysis", {})
    add(f"- **Primary Country :** {geo.get('primary_country', '—')} ({geo.get('country_code', '?')})")
    add(f"- **Flagged Region  :** {'Yes ⚠' if geo.get('is_flagged_region') else 'No'}")
    if geo.get("risk_note"):
        add(f"- **Region Note     :** {geo.get('risk_note')}")
    add(f"- **Geo Source      :** {geo.get('geo_source', 'WHOIS correlation')}")
    threat_region = data.get("threat_region", {})
    add(f"- **Hosting Country :** {threat_region.get('hosting_country', '—')}")
    add(f"- **Region Risk     :** {threat_region.get('risk_level', '—')}")
    add("")

    # IP breakdown table
    ip_geo = geo.get("ip_geo", [])
    if ip_geo:
        add("| IP Address | Country | Flagged |")
        add("|------------|---------|---------|")
        for entry in ip_geo:
            add(f"| {entry.get('ip','?')} | {entry.get('country','?')} | "
                f"{'⚠ Yes' if entry.get('flagged') else 'No'} |")
        add("")

    # ── Section 7: Advanced Threat Modules ───────────────────────
    add(_section("7. Advanced Threat Modules"))

    # Look-alike
    lookalike = data.get("lookalike", {})
    add("### 7.1  Anti-Phishing — Look-Alike Domain Detection")
    if lookalike and lookalike.get("is_lookalike"):
        add(f"- **Match          :** {lookalike.get('matched_domain', '—')}")
        add(f"- **Similarity     :** {lookalike.get('similarity_score', 0)}%")
        add(f"- **Classification :** {lookalike.get('risk_classification', '—')}")
    else:
        add("- No look-alike domain pattern detected.")
    add("")

    # Phishing Simulation
    phish = data.get("phish_sim", {})
    add("### 7.2  Phishing Simulation Engine")
    add(f"- **Harvest Probability  :** {phish.get('phishing_probability', 0)}%")
    add(f"- **Behavior Class       :** {phish.get('behavior_classification', '—')}")
    add("")

    # Domain DNA
    dna = data.get("domain_dna", {})
    add("### 7.3  Domain DNA Profile")
    add(f"- **Structure Type    :** {dna.get('structure_type', '—')}")
    add(f"- **Entropy Level     :** {dna.get('entropy_level', '—')}")
    add(f"- **Behavioral Pattern:** {dna.get('behavioral_pattern', '—')}")
    add(f"- **DNA Summary Score :** {dna.get('summary_score', '—')} / 10")
    add("")

    # Threat Alerts
    alerts = data.get("threat_alerts", [])
    add("### 7.4  Active Threat Alerts")
    if alerts:
        for al in alerts:
            add(f"  - **{al.get('level','?')} :** {al.get('msg','')}")
    else:
        add("  - No active threat alerts.")
    add("")

    # URL Pattern
    url_info = data.get("url_analysis", {})
    add("### 7.5  URL Pattern Analysis")
    add(f"- **URL Length        :** {url_info.get('url_length', '—')} chars")
    add(f"- **TLD               :** {url_info.get('tld', '—')} "
        f"{'⚠ Suspicious' if url_info.get('suspicious_tld') else ''}")
    add(f"- **Subdomain Depth   :** {url_info.get('subdomain_depth', 0)}")
    add(f"- **Phishing Score    :** {url_info.get('phishing_score', 0)} / 100")
    url_flags = url_info.get("flags", [])
    if url_flags:
        add(f"- **Flags             :** {'; '.join(url_flags)}")
    add("")

    # ── Section 8: Attack Surface Summary ────────────────────────
    add(_section("8. Attack Surface Summary"))
    atk = data.get("attack_surface", {})
    for key, label in [
        ("proxy_layer",      "Proxy Layer"),
        ("dns_strength",     "DNS Strength"),
        ("threat_footprint", "Threat Footprint"),
        ("exposure_level",   "Exposure Level"),
    ]:
        layer = atk.get(key, {})
        add(f"- **{label}** : {layer.get('status','—')} — {layer.get('detail','—')} "
            f"[{layer.get('risk','?')} risk]")
    add("")

    # ── Section 9: Investigation Status & Classification ─────────
    add(_section("9. Investigation Status"))
    inv = data.get("investigation_status", {})
    add(f"- **Visibility          :** {inv.get('visibility', '—')}")
    add(f"- **Escalation Required :** {'Yes ⚠' if inv.get('escalation_required') else 'No'}")
    add(f"- **Escalation Note     :** {inv.get('escalation_note', '—')}")
    add("")

    cls = data.get("classification", {})
    add(f"- **Verdict             :** {cls.get('verdict', '—')}")
    add(f"- **TLP                 :** {cls.get('tlp', 'TLP:WHITE')}")
    add(f"- **Retention           :** {cls.get('retention', 'Standard 90-day')}")
    tags = cls.get("tags", [])
    if tags:
        add(f"- **Tags                :** {', '.join(tags)}")
    add("")

    # ── Section 10: AI Model Features ────────────────────────────
    features = data.get("features", {})
    if features:
        add(_section("10. AI Model Feature Extraction"))
        add("| Feature | Raw Value |")
        add("|---------|-----------|")
        for k, v in features.items():
            add(f"| {k} | {v} |")
        add("")

    # ── Section 11: Chart Descriptions ───────────────────────────
    add(_section("11. Chart Descriptions"))
    add("The following charts are rendered in the PDF version of this report:")
    add("")
    add("**A. Risk Score Gauge**")
    add(f"  Semicircular gauge displaying the current risk score of {risk_score}/100 "
        f"in the {severity} zone. Red = High (≥71), Yellow = Medium (31–70), Green = Low (≤30).")
    add("")
    add("**B. Redirect Chain Flowchart**")
    n_hops = rc.get("total_hops", 0)
    if n_hops > 1:
        add(f"  The target URL traversed {n_hops} hop(s) before reaching the final endpoint. "
            f"{'This is SUSPICIOUS — redirect count exceeds 2.' if suspicious_chain else 'Chain length is within normal range.'}")
    else:
        add("  No redirect hops detected. Single-hop direct connection.")
    add("")
    add("**C. Risk Trend History**")
    add("  Line chart showing how risk scores have changed across all scans performed "
        "in this browser session. Helps identify patterns or anomalies across multiple targets.")
    add("")

    # ── Footer ────────────────────────────────────────────────────
    add(_divider("═"))
    add(f"**Generated by:** HostTrace AI v8.0 — Advanced Cyber Intelligence Platform")
    add(f"**Classification:** {cls.get('tlp', 'TLP:WHITE')} | For authorized security research only")
    add(f"**Timestamp:** {scan_dt}")
    add(_divider("═"))

    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════
# STEP 2 — PDF RENDERER  (ReportLab layout engine)
# ══════════════════════════════════════════════════════════════════

def render_pdf_from_text(md_text: str, data: dict) -> bytes:
    """
    Takes the Markdown text report and renders it into a styled PDF.
    Embeds matplotlib charts (risk gauge, redirect chart, risk trend).
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    # ── Styles ───────────────────────────────────────────────────
    styles = getSampleStyleSheet()
    DARK_BLUE  = colors.HexColor("#0f3d5e")
    MED_BLUE   = colors.HexColor("#2a5e82")
    LIGHT_GRAY = colors.HexColor("#f4f7f8")
    BORDER     = colors.HexColor("#aabfd1")

    t_title = ParagraphStyle("T_Title", parent=styles["h1"],
        fontName=FONT_NAME, fontSize=19, textColor=DARK_BLUE,
        spaceAfter=10, spaceBefore=0)
    t_h2 = ParagraphStyle("T_H2", parent=styles["h2"],
        fontName=FONT_NAME, fontSize=13, textColor=MED_BLUE,
        spaceBefore=14, spaceAfter=5)
    t_h3 = ParagraphStyle("T_H3", parent=styles["h3"],
        fontName=FONT_NAME, fontSize=11, textColor=MED_BLUE,
        spaceBefore=8, spaceAfter=4)
    t_body = ParagraphStyle("T_Body", parent=styles["Normal"],
        fontName=FONT_NAME, fontSize=9.5, leading=14, spaceAfter=3)
    t_mono = ParagraphStyle("T_Mono", parent=styles["Normal"],
        fontName="Courier", fontSize=8.5, leading=12, spaceAfter=2,
        leftIndent=10, textColor=colors.HexColor("#1e3a5f"))
    t_foot = ParagraphStyle("T_Foot", parent=styles["Normal"],
        fontName=FONT_NAME, fontSize=7.5, textColor=colors.gray, alignment=1)
    t_sev_high = ParagraphStyle("T_SevH", parent=t_body,
        textColor=colors.HexColor("#ff4f6d"), fontName=FONT_NAME, fontsize=10)
    t_sev_med  = ParagraphStyle("T_SevM", parent=t_body,
        textColor=colors.HexColor("#ffd166"), fontName=FONT_NAME)
    t_sev_low  = ParagraphStyle("T_SevL", parent=t_body,
        textColor=colors.HexColor("#00c97a"), fontName=FONT_NAME)

    story = []

    def add(text, style=t_body):
        """Safely add a paragraph, converting markdown to ReportLab XML."""
        try:
            story.append(Paragraph(_p(str(text)), style))
        except Exception as e:
            # Last-resort fallback: strip everything to plain ASCII
            plain = str(text).encode("ascii", "ignore").decode("ascii")
            try:
                story.append(Paragraph(plain, style))
            except Exception:
                pass  # Skip unrenderable lines rather than crash

    def hr():
        story.append(HRFlowable(width="100%", thickness=0.8,
                                color=BORDER, spaceAfter=8, spaceBefore=4))

    def sp(h=8):
        story.append(Spacer(1, h))

    def table(rows, col_widths, header_row=True):
        if not rows:
            return
        t = Table(rows, colWidths=col_widths)
        ts = [
            ("FONTNAME", (0, 0), (-1, -1), FONT_NAME),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("ALIGN",    (0, 0), (-1, -1), "LEFT"),
            ("VALIGN",   (0, 0), (-1, -1), "MIDDLE"),
            ("GRID",     (0, 0), (-1, -1), 0.5, BORDER),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [LIGHT_GRAY, colors.white]),
        ]
        if header_row:
            ts += [
                ("BACKGROUND", (0, 0), (-1, 0), MED_BLUE),
                ("TEXTCOLOR",  (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTSIZE",   (0, 0), (-1, 0), 9),
            ]
        t.setStyle(TableStyle(ts))
        story.append(t)
        sp(6)

    def embed_image(buf, width_pts, height_pts, caption=""):
        if buf:
            story.append(Image(buf, width=width_pts, height=height_pts))
            if caption:
                try:
                    story.append(Paragraph(_p(caption), t_foot))
                except Exception:
                    pass
            sp(8)

    # ── Page 1: Header ───────────────────────────────────────────
    add("HostTrace AI — Forensic Intelligence Report", t_title)
    hr()

    scan_dt   = data.get("scan_timestamp", "—")
    domain    = data.get("domain", "Unknown")
    input_url = data.get("input_url", domain)
    risk_sc   = data.get("risk_score", 0)
    risk_lv   = data.get("risk_level", "Unknown")

    meta_rows = [
        ["Scan Timestamp", scan_dt],
        ["Report ID",      data.get("report_id", "—")],
        ["Trace ID",       data.get("trace_id", "—")],
        ["Input URL/IP",   input_url],
        ["Canonical Domain", domain],
        ["IP Address(es)", ", ".join(data.get("ip_addresses", [])) or "Unresolved"],
    ]
    table([[_p(k), _p(str(v))] for k, v in meta_rows],
          col_widths=[120, 380], header_row=False)

    # ── Risk gauge chart ─────────────────────────────────────────
    gauge_buf = _chart_risk_gauge(risk_sc, risk_lv)
    embed_image(gauge_buf, 280, 150, f"Fig A — Risk Score Gauge: {risk_sc}/100 ({risk_lv})")

    # ── Section headings derived from Markdown ───────────────────
    # Parse the already-built Markdown text and render it into ReportLab
    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            sp(4)
            continue

        # Skip already-rendered title
        if stripped.startswith("# ") and "HostTrace AI" in stripped:
            continue
        # Skip decorative dividers (═══ lines and ## section markers)
        if (stripped.startswith("="*3) or stripped.startswith("\u2550"*3)
                or stripped.startswith("\u2500"*3) or stripped.startswith("---")):
            hr()
            continue
        if stripped.startswith("## ") or stripped.startswith("### "):
            # Also skip --- separator lines underneath them
            continue
        # Headings (from _section())
        if stripped.startswith("# "):
            add(stripped[2:], t_h2)
        # Table separator row  |---|----|  -> skip
        elif stripped.startswith("|") and re.match(r'^\|[-| :]+\|$', stripped):
            continue
        # Markdown table data row
        elif stripped.startswith("|") and stripped.endswith("|"):
            add(stripped, t_mono)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("  - "):
            add(stripped, t_body)
        # Block quote
        elif stripped.startswith("> "):
            add(stripped[2:], t_mono)
        else:
            add(stripped, t_body)

    # ── Charts section ───────────────────────────────────────────
    sp(12)
    hr()
    add("Embedded Visualizations", t_h2)

    rc_chain = data.get("redirect_chain", {}).get("chain", [])
    if len(rc_chain) > 1:
        rc_buf = _chart_redirect_chain(rc_chain)
        embed_image(rc_buf, 430, max(120, len(rc_chain) * 58),
                    "Fig B — Redirect Chain Flowchart")

    scan_history = data.get("scan_history", [])
    if scan_history and len(scan_history) >= 2:
        trend_buf = _chart_risk_trend(scan_history)
        embed_image(trend_buf, 430, 200,
                    "Fig C — Risk Trend History (session scans)")
    else:
        add("Note: Risk Trend chart requires ≥2 scans in the current session. "
            "Scan additional domains to generate the trend view.", t_body)

    # ── Footer ───────────────────────────────────────────────────
    sp(20)
    hr()
    add("Generated by HostTrace AI v8.0 — Advanced Cyber Intelligence Platform", t_foot)
    add(f"Classification: {data.get('classification', {}).get('tlp', 'TLP:WHITE')} | "
        "For authorized security research only", t_foot)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ══════════════════════════════════════════════════════════════════
# STEP 3 — CONVENIENCE WRAPPER (kept for backward compat)
# ══════════════════════════════════════════════════════════════════

def generate_pdf_report(data: dict) -> bytes:
    """
    Convenience function: builds the text report first, then renders it to PDF.
    """
    md_text = generate_text_report(data)
    return render_pdf_from_text(md_text, data)


# ══════════════════════════════════════════════════════════════════
# STEP 4 — WORD (.docx) REPORT GENERATOR
# ══════════════════════════════════════════════════════════════════

def generate_word_report(data: dict) -> bytes:
    """
    Generates a styled Word (.docx) forensic report from scan data.
    Uses python-docx to build a fully formatted document with:
      - Cover metadata table
      - All scan sections with headings & bullet lists
      - Embedded Matplotlib charts as inline images
      - Professional color scheme & styles
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re as _re

    # ── Helper: safe ASCII text ──────────────────────────────────
    def _w(text) -> str:
        """Convert any value to safe ASCII string for Word."""
        s = str(text)
        for uni, asc in SPECIAL_CHARS.items():
            s = s.replace(uni, asc)
        return s.encode("ascii", "ignore").decode("ascii")

    # ── Color palette ────────────────────────────────────────────
    C_DARK   = RGBColor(0x0f, 0x3d, 0x5e)   # dark navy
    C_MED    = RGBColor(0x2a, 0x5e, 0x82)   # medium blue
    C_ACCENT = RGBColor(0x00, 0xc9, 0x7a)   # green
    C_RED    = RGBColor(0xff, 0x4f, 0x6d)   # red/danger
    C_YELLOW = RGBColor(0xff, 0xd1, 0x66)   # yellow/warning
    C_WHITE  = RGBColor(0xff, 0xff, 0xff)
    C_LGRAY  = RGBColor(0xf4, 0xf7, 0xf8)

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Style helpers ─────────────────────────────────────────────
    def set_cell_bg(cell, hex_color: str):
        """Set table cell background colour."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"),  "clear")
        tcPr.append(shd)

    def add_heading(text: str, level: int = 1,
                    color: RGBColor = C_DARK) -> None:
        p = doc.add_heading(_w(text), level=level)
        run = p.runs[0] if p.runs else p.add_run(_w(text))
        run.font.color.rgb = color
        run.font.bold = True
        p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
        p.paragraph_format.space_after  = Pt(4)

    def add_para(text: str, bold_pattern: bool = True,
                 indent: int = 0, italic: bool = False) -> None:
        """Add a paragraph, rendering **bold** segments."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent * 0.5)
        p.paragraph_format.space_after = Pt(2)
        raw = _w(text)
        if bold_pattern:
            parts = _re.split(r'(\*\*[^*]+\*\*)', raw)
            for part in parts:
                m = _re.match(r'\*\*([^*]+)\*\*', part)
                if m:
                    run = p.add_run(m.group(1))
                    run.bold = True
                    run.font.color.rgb = C_DARK
                else:
                    run = p.add_run(part)
                    run.italic = italic
                    run.font.size = Pt(10)
        else:
            run = p.add_run(raw)
            run.italic = italic
            run.font.size = Pt(10)

    def add_bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        raw = _w(text.lstrip("- ").lstrip("  - "))
        parts = _re.split(r'(\*\*[^*]+\*\*)', raw)
        for part in parts:
            m = _re.match(r'\*\*([^*]+)\*\*', part)
            if m:
                run = p.add_run(m.group(1))
                run.bold = True
            else:
                run = p.add_run(part)
        p.paragraph_format.space_after = Pt(1)

    def add_rule():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),  "single")
        bottom.set(qn("w:sz"),   "6")
        bottom.set(qn("w:space"),"1")
        bottom.set(qn("w:color"),"2a5e82")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def meta_table(rows: list) -> None:
        """Two-column key-value table for cover metadata."""
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.columns[0].width = Cm(4.5)
        t.columns[1].width = Cm(13)
        for i, (k, v) in enumerate(rows):
            k_cell = t.cell(i, 0)
            v_cell = t.cell(i, 1)
            set_cell_bg(k_cell, "DCE9F5")
            k_run = k_cell.paragraphs[0].add_run(_w(k))
            k_run.bold = True
            k_run.font.color.rgb = C_DARK
            k_run.font.size = Pt(9)
            v_run = v_cell.paragraphs[0].add_run(_w(str(v)))
            v_run.font.size = Pt(9)
        doc.add_paragraph()

    def two_col_table(rows: list, headers: list = None) -> None:
        """Generic two-column data table with optional header row."""
        total = len(rows) + (1 if headers else 0)
        t = doc.add_table(rows=total, cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.columns[0].width = Cm(7)
        t.columns[1].width = Cm(10.5)
        row_offset = 0
        if headers:
            for ci, h in enumerate(headers[:2]):
                cell = t.cell(0, ci)
                set_cell_bg(cell, "2a5e82")
                run = cell.paragraphs[0].add_run(_w(h))
                run.bold = True
                run.font.color.rgb = C_WHITE
                run.font.size = Pt(9)
            row_offset = 1
        for ri, (k, v) in enumerate(rows):
            for ci, val in enumerate([k, v]):
                cell = t.cell(ri + row_offset, ci)
                if (ri + row_offset) % 2 == 1:
                    set_cell_bg(cell, "F4F7F8")
                run = cell.paragraphs[0].add_run(_w(str(val)))
                run.font.size = Pt(9)
        doc.add_paragraph()

    def embed_chart(buf, width_inches: float = 5.5, caption: str = ""):
        """Embed a BytesIO PNG chart into the document."""
        if not buf:
            return
        try:
            buf.seek(0)
            doc.add_picture(buf, width=Inches(width_inches))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cp = doc.add_paragraph(_w(caption))
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.runs[0].italic = True
                cp.runs[0].font.size = Pt(8)
                cp.runs[0].font.color.rgb = C_MED
            doc.add_paragraph()
        except Exception as e:
            print(f"[Word] Image embed error: {e}")

    # ── Retrieve key data ─────────────────────────────────────────
    scan_dt   = _w(data.get("scan_timestamp", "—"))
    domain    = _w(data.get("domain", "Unknown"))
    input_url = _w(data.get("input_url", domain))
    risk_sc   = data.get("risk_score", 0)
    risk_lv   = _w(data.get("risk_level", "Unknown"))
    verdict   = data.get("verdict", {})
    pred      = data.get("ai_prediction", {})
    whois     = data.get("whois", {})
    ssl       = data.get("ssl_analysis", {})
    http_a    = data.get("http_analysis", {})
    rc        = data.get("redirect_chain", {})
    geo       = data.get("geo_analysis", {})
    reg_cls   = data.get("classification", {})

    sev_color = C_RED if risk_sc >= 71 else (C_YELLOW if risk_sc >= 31 else C_ACCENT)
    severity  = "HIGH" if risk_sc >= 71 else ("MEDIUM" if risk_sc >= 31 else "LOW")

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("HostTrace AI — Forensic Intelligence Report")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = C_DARK

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Advanced Cybersecurity Investigation Platform  |  v8.0")
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.color.rgb = C_MED

    add_rule()

    # Risk badge paragraph
    badge_p = doc.add_paragraph()
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = badge_p.add_run(f"  RISK: {severity}  |  Score: {risk_sc}/100  |  {_w(verdict.get('status', '—'))}  ")
    br.bold = True
    br.font.size = Pt(12)
    br.font.color.rgb = sev_color
    doc.add_paragraph()

    # Risk gauge chart
    gauge_buf = _chart_risk_gauge(risk_sc, risk_lv)
    embed_chart(gauge_buf, width_inches=3.5, caption=f"Risk Score Gauge: {risk_sc}/100 ({risk_lv})")

    # Cover metadata
    meta_table([
        ("Scan Timestamp",   scan_dt),
        ("Report ID",        data.get("report_id", "—")),
        ("Trace ID",         data.get("trace_id", "—")),
        ("Input URL / IP",   input_url),
        ("Canonical Domain", domain),
        ("IP Address(es)",   ", ".join(data.get("ip_addresses", [])) or "Unresolved"),
        ("Proxy / CDN",      f"Detected ({data.get('proxy_provider','')})" if data.get("proxy_detected") else "Not Detected"),
        ("Infrastructure",   data.get("real_hosting", "—")),
        ("AI Confidence",    f"{verdict.get('confidence', '—')}%"),
        ("Classification",   reg_cls.get("tlp", "TLP:WHITE")),
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 1: AI RISK PREDICTION
    # ══════════════════════════════════════════════════════════════
    add_heading("1. AI Risk Prediction", level=1)
    add_rule()

    add_para(f"**Final Verdict :** {_w(verdict.get('status', '—'))}")
    add_para(f"**Risk Score    :** {risk_sc} / 100  ({severity} severity)")
    add_para(f"**AI Confidence :** {verdict.get('confidence', '—')}%")
    if pred:
        add_para(f"**ML Model      :** RandomForestClassifier  |  Accuracy: {pred.get('accuracy', 91.6):.1f}%")
        add_para(f"**Prediction    :** {_w(pred.get('prediction', 'UNKNOWN'))}")
        probs = pred.get("probabilities", {})
        add_para(f"**Probabilities :** SAFE={probs.get('SAFE',0):.1f}%  |  "
                 f"SUSPICIOUS={probs.get('SUSPICIOUS',0):.1f}%  |  "
                 f"DANGEROUS={probs.get('DANGEROUS',0):.1f}%")

    risk_reason = data.get("risk_reason", {})
    if risk_reason and isinstance(risk_reason, dict):
        doc.add_paragraph()
        add_heading("Risk Explanation (XAI)", level=2, color=C_MED)
        add_para(f"**{_w(risk_reason.get('title', ''))}**")
        for pt in risk_reason.get("points", []):
            add_bullet(f"- {pt}")

    # Risk breakdown table
    rb = data.get("risk_breakdown", {})
    if rb:
        doc.add_paragraph()
        add_heading("Risk Factor Breakdown", level=2, color=C_MED)
        rb_rows = [(k.replace("_", " ").title(), f"+{v}" if v > 0 else str(v))
                   for k, v in rb.items() if v != 0]
        if rb_rows:
            two_col_table(rb_rows, headers=["Component", "Score Impact"])

    # XAI signals
    xai = data.get("xai_signals", [])
    if xai:
        doc.add_paragraph()
        add_heading("Explainable AI (XAI) Signals", level=2, color=C_MED)
        for sig in xai:
            add_bullet(f"- **{sig.get('feature','?')} [{sig.get('impact','?')}]:** {sig.get('desc','')}")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 2: OSINT ANALYSIS
    # ══════════════════════════════════════════════════════════════
    add_heading("2. OSINT Analysis", level=1)
    add_rule()

    add_heading("2.1  WHOIS Forensics", level=2, color=C_MED)
    two_col_table([
        ("Registrar",    whois.get("registrar", "—")),
        ("Organisation", whois.get("org", "—")),
        ("Country",      whois.get("country", "—")),
        ("Created",      whois.get("creation_date", "—")),
        ("Expires",      whois.get("expiry_date", "—")),
        ("DNSSEC",       whois.get("dnssec", "—")),
        ("Name Servers", ", ".join(whois.get("name_servers", [])) or "—"),
    ], headers=["Field", "Value"])

    add_heading("2.2  DNS & Hosting", level=2, color=C_MED)
    two_col_table([
        ("Resolved IPs",       ", ".join(data.get("ip_addresses", [])) or "None"),
        ("Detected Hosting",   data.get("real_hosting", "—")),
        ("Possible Hosting",   data.get("possible_hosting") or "—"),
        ("Hosting Confidence", f"{data.get('confidence', '—')}%"),
        ("CDN Detected",       f"Yes — {data.get('cdn_provider','')}" if data.get("cdn_detected") else "No"),
        ("WAF Suspected",      "Yes" if data.get("waf_suspected") else "No"),
        ("Masking Level",      data.get("masking_level", "—")),
    ], headers=["Field", "Value"])

    # OSINT historical records
    osint = data.get("osint_simulation", {})
    if osint:
        add_heading("2.3  Historical Domain Reputation", level=2, color=C_MED)
        add_para(f"- {_w(osint.get('exposure_note', '—'))}")
        recs = osint.get("records", [])
        if recs:
            two_col_table(
                [(f"{r.get('date','?')} — {r.get('ip','?')}",
                  f"{r.get('provider','?')}  [{r.get('type','?')}]") for r in recs],
                headers=["Date / IP", "Provider / Type"]
            )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 3: THREAT INTELLIGENCE
    # ══════════════════════════════════════════════════════════════
    add_heading("3. Threat Intelligence", level=1)
    add_rule()

    threat = data.get("threat_intel", {})
    two_col_table([
        ("VirusTotal Flags",  threat.get("virustotal_flags", 0)),
        ("Blacklist Hits",    threat.get("blacklist_hits", 0)),
        ("Abuse.ch C2",       "Yes [!]" if threat.get("abuse_ch") else "No"),
        ("OTX Pulses",        threat.get("otx_pulses", 0)),
        ("C2 Detected",       "Yes [!]" if threat.get("c2_detected") else "No"),
        ("Phishing Category", "Yes [!]" if threat.get("phishing_category") else "No"),
    ], headers=["Indicator", "Result"])

    add_heading("3.1  SSL / TLS Certificate", level=2, color=C_MED)
    two_col_table([
        ("Valid",       "Yes" if ssl.get("ssl_valid") else "No"),
        ("Expired",     "Yes [!]" if ssl.get("ssl_expired") else "No"),
        ("Self-Signed", "Yes [!]" if ssl.get("self_signed") else "No"),
        ("Grade",       ssl.get("grade", "N/A")),
        ("Issuer",      str(ssl.get("issuer", {}).get("organizationName", ssl.get("issuer", "—")))),
        ("Valid From",  ssl.get("valid_from", "—")),
        ("Valid Until", ssl.get("valid_until", "—")),
        ("Days Left",   ssl.get("days_remaining", "—")),
    ], headers=["Field", "Value"])

    add_heading("3.2  HTTP Security Headers", level=2, color=C_MED)
    two_col_table([
        ("Server",         http_a.get("server", "Unknown")),
        ("Security Score", f"{http_a.get('security_score', 0)}%"),
        ("HSTS",           "Enabled" if http_a.get("hsts") else "Missing"),
        ("CSP",            "Present" if http_a.get("csp") else "Missing"),
        ("X-Frame-Options","Present" if http_a.get("x_frame") else "Missing"),
        ("Present Headers",", ".join(http_a.get("present_security_headers", [])) or "None"),
        ("Missing Headers",", ".join(http_a.get("missing_security_headers", [])) or "None"),
    ], headers=["Header", "Status"])

    # Threat alerts
    alerts = data.get("threat_alerts", [])
    if alerts:
        add_heading("3.3  Active Threat Alerts", level=2, color=C_MED)
        for al in alerts:
            add_bullet(f"- [{al.get('level','?')}] {al.get('msg','')}")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 4: REDIRECT CHAIN
    # ══════════════════════════════════════════════════════════════
    add_heading("4. Redirect Chain Analysis", level=1)
    add_rule()

    add_para(f"**Total Hops     :** {rc.get('total_hops', 0)}")
    add_para(f"**Redirect Count :** {rc.get('redirect_count', 0)}")
    add_para(f"**Suspicious     :** {'Yes [!] — Long chain detected' if rc.get('suspicious') else 'No'}")
    doc.add_paragraph()

    chain = rc.get("chain", [])
    if chain:
        two_col_table(
            [(f"Hop {h.get('hop','?')}  [{h.get('status','?')}]",
              _w(str(h.get('url','—'))[:80]) + ("  [FINAL]" if h.get("final") else "  -> redirect"))
             for h in chain],
            headers=["Hop / Status", "URL"]
        )
        # Embed redirect chain chart
        rc_buf = _chart_redirect_chain(chain) if len(chain) > 1 else None
        embed_chart(rc_buf, width_inches=5.5, caption="Fig B — Redirect Chain Flowchart")
    else:
        add_para("No redirect hops recorded. Direct connection established.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 5: GEOLOCATION & ASN
    # ══════════════════════════════════════════════════════════════
    add_heading("5. Geolocation & ASN Intelligence", level=1)
    add_rule()

    threat_region = data.get("threat_region", {})
    two_col_table([
        ("Primary Country",  f"{geo.get('primary_country','—')} ({geo.get('country_code','?')})"),
        ("Flagged Region",   "Yes [!]" if geo.get("is_flagged_region") else "No"),
        ("Region Risk Note", geo.get("risk_note", "—") or "None"),
        ("Hosting Country",  threat_region.get("hosting_country", "—")),
        ("Region Risk",      threat_region.get("risk_level", "—")),
    ], headers=["Field", "Value"])

    asn = data.get("asn_analysis", {})
    if asn:
        two_col_table([
            ("Proxy ASN",    asn.get("proxy_provider", "—")),
            ("Origin ASN",   asn.get("origin_asn_provider", "—")),
            ("ASN Mismatch", "Yes [!] — " + asn.get("mismatch_note", "") if asn.get("mismatch_detected") else "No"),
        ], headers=["ASN Field", "Value"])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 6: ADVANCED THREAT MODULES
    # ══════════════════════════════════════════════════════════════
    add_heading("6. Advanced Threat Modules", level=1)
    add_rule()

    lookalike = data.get("lookalike", {})
    add_heading("6.1  Anti-Phishing — Look-Alike Domain", level=2, color=C_MED)
    if lookalike and lookalike.get("is_lookalike"):
        two_col_table([
            ("Matched Domain", lookalike.get("matched_domain", "—")),
            ("Similarity",     f"{lookalike.get('similarity_score', 0)}%"),
            ("Classification", lookalike.get("risk_classification", "—")),
        ], headers=["Field", "Value"])
    else:
        add_para("No look-alike domain pattern detected.")

    phish = data.get("phish_sim", {})
    add_heading("6.2  Phishing Simulation Engine", level=2, color=C_MED)
    two_col_table([
        ("Harvest Probability", f"{phish.get('phishing_probability', 0)}%"),
        ("Behavior Class",      phish.get("behavior_classification", "—")),
    ], headers=["Field", "Value"])

    dna = data.get("domain_dna", {})
    add_heading("6.3  Domain DNA Profile", level=2, color=C_MED)
    two_col_table([
        ("Structure Type",     dna.get("structure_type", "—")),
        ("Entropy Level",      dna.get("entropy_level", "—")),
        ("Behavioral Pattern", dna.get("behavioral_pattern", "—")),
        ("DNA Score",          f"{dna.get('summary_score','—')} / 10"),
    ], headers=["Field", "Value"])

    url_info = data.get("url_analysis", {})
    add_heading("6.4  URL Pattern Analysis", level=2, color=C_MED)
    two_col_table([
        ("URL Length",      f"{url_info.get('url_length','—')} chars"),
        ("TLD",             f"{url_info.get('tld','—')}{'  [!] Suspicious' if url_info.get('suspicious_tld') else ''}"),
        ("Subdomain Depth", url_info.get("subdomain_depth", 0)),
        ("Phishing Score",  f"{url_info.get('phishing_score',0)} / 100"),
        ("Flags",           "; ".join(url_info.get("flags", [])) or "None"),
    ], headers=["Field", "Value"])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 7: RISK TREND CHART
    # ══════════════════════════════════════════════════════════════
    scan_history = data.get("scan_history", [])
    if scan_history and len(scan_history) >= 2:
        add_heading("7. Risk Trend History", level=1)
        add_rule()
        trend_buf = _chart_risk_trend(scan_history)
        embed_chart(trend_buf, width_inches=5.5, caption="Fig C — Risk Trend History (session scans)")
        two_col_table(
            [(f"#{i+1}  {h.get('domain','?')[:25]}", f"{h.get('risk',0)}  ({h.get('pred','?')})")
             for i, h in enumerate(scan_history)],
            headers=["Scan", "Risk Score (Prediction)"]
        )

    # ══════════════════════════════════════════════════════════════
    # SECTION 8: AI FEATURE EXTRACTION
    # ══════════════════════════════════════════════════════════════
    features = data.get("features", {})
    if features:
        add_heading("8. AI Model Feature Extraction", level=1)
        add_rule()
        two_col_table(
            [(str(k), str(v)) for k, v in features.items()],
            headers=["Feature Node", "Raw Value"]
        )

    # ══════════════════════════════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════════════════════════════
    add_rule()
    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot_p.add_run(
        f"Generated by HostTrace AI v8.0  |  "
        f"{reg_cls.get('tlp', 'TLP:WHITE')}  |  "
        f"For authorized security research only  |  {scan_dt}"
    )
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = C_MED

    # ── Serialize to bytes ────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
